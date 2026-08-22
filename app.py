"""
Q&A Evaluation App
------------------
Upload a baseline (reference) Question + Answer, then upload one or more
individual answer files (PDF or DOCX). The app extracts each individual's
question/answer using an LLM (Groq), scores each answer against the
baseline, and ranks all candidates.
"""

import os
import io
import json
import time
from pathlib import Path

import pandas as pd
import streamlit as st
from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel
from pypdf import PdfReader
from docx import Document

try:
    from sentence_transformers import SentenceTransformer, util as st_util
    SEMANTIC_MODEL_AVAILABLE = True
except ImportError:
    SEMANTIC_MODEL_AVAILABLE = False


# ============================================================
# PAGE CONFIG
# ============================================================

st.set_page_config(
    page_title="Q&A Evaluator",
    page_icon="📝",
    layout="wide",
)

load_dotenv()


# ============================================================
# DATA MODELS
# ============================================================

class QA(BaseModel):
    Question: str
    Answer: str


class IndividualAns(BaseModel):
    Ques: str
    Answer: str


class MatchResult(BaseModel):
    score: float
    details: str


# ============================================================
# FILE READING HELPERS
# ============================================================

def read_pdf(uploaded_file) -> str:
    """Extract text from an in-memory PDF (Streamlit UploadedFile)."""
    reader = PdfReader(uploaded_file)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text


def read_docx(uploaded_file) -> str:
    """Extract text (paragraphs + tables) from an in-memory DOCX file."""
    document = Document(uploaded_file)
    text = ""

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"

    return text


def read_uploaded_file(uploaded_file) -> str:
    """Dispatch to the correct reader based on file extension."""
    extension = Path(uploaded_file.name).suffix.lower()

    # Reset pointer in case the file was already read once
    uploaded_file.seek(0)

    if extension == ".pdf":
        return read_pdf(uploaded_file)
    elif extension == ".docx":
        return read_docx(uploaded_file)
    else:
        return ""


# ============================================================
# LLM HELPERS
# ============================================================

def get_client(api_key: str) -> Groq:
    return Groq(api_key=api_key)


@st.cache_resource(show_spinner=False)
def load_semantic_model():
    """Load once and cache across reruns/files."""
    return SentenceTransformer("all-MiniLM-L6-v2")


def semantic_similarity(model, text_a: str, text_b: str) -> float:
    """Cosine similarity between two texts' embeddings, scaled to 0-100.

    This captures meaning-based closeness even when the wording differs
    a lot (paraphrases, synonyms, reordered points, etc.) — unlike exact
    or fuzzy string matching, which only looks at surface-level overlap.
    """
    if not text_a.strip() or not text_b.strip():
        return 0.0

    embeddings = model.encode([text_a, text_b], convert_to_tensor=True)
    cosine_sim = st_util.cos_sim(embeddings[0], embeddings[1]).item()

    # cos_sim ranges roughly -1..1; clamp then scale to 0..100
    cosine_sim = max(0.0, min(1.0, cosine_sim))
    return round(cosine_sim * 100, 2)


def parse_QA(client: Groq, model: str, QA_text: str) -> IndividualAns:
    """Ask the LLM to extract a Question + Answer pair from raw document text."""
    schema = IndividualAns.model_json_schema()

    system_prompt = f"""
You are an expert Question Answer parser.

Your job is to extract the question and the individual's answer
from the provided document.

Return ONLY valid JSON.

The response must contain the word JSON.

Return JSON matching this schema:

{schema}

IMPORTANT RULES:

1. Do not invent information.
2. Extract the actual question from the document.
3. Extract the individual's actual answer.
4. Keep the answer as close to the original meaning as possible.
5. If the question cannot be found, return an empty string.
6. If the answer cannot be found, return an empty string.
"""

    user_prompt = f"""
Analyze the following Question Answer document:

{QA_text}
"""

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ]

    response = client.chat.completions.create(
        model=model,
        messages=messages,
        response_format={"type": "json_object"},
    )

    raw_output = response.choices[0].message.content
    data = json.loads(raw_output)
    return IndividualAns(**data)


def final_score(client: Groq, model: str, correct_answer: QA, individual_answer: IndividualAns) -> MatchResult:
    """Ask the LLM to score an individual answer against the reference answer."""
    match_schema = MatchResult.model_json_schema()

    prompt = f"""
You are an expert answer evaluator.

Your job is to compare an individual's answer with the
correct/reference answer.

IMPORTANT: Grade based on SEMANTIC MEANING, not exact wording.

- The individual's answer does NOT need to match the reference
  word-for-word, use the same sentence structure, or list points
  in the same order.
- Paraphrases, synonyms, and differently-worded statements that
  convey the same fact or idea should be treated as fully correct.
- Only reduce the score when the individual's answer is missing
  information present in the reference, adds information that
  contradicts the reference, or is factually/conceptually different
  from what the reference conveys.
- Do not penalize differences in phrasing, length, tone, or ordering
  as long as the underlying meaning is preserved.

REFERENCE QUESTION:

{correct_answer.Question}


REFERENCE/CORRECT ANSWER:

{correct_answer.Answer}


INDIVIDUAL QUESTION:

{individual_answer.Ques}


INDIVIDUAL ANSWER:

{individual_answer.Answer}


Evaluate the individual's answer against the correct answer.

Give a score from 0 to 100.

Scoring guidelines:

90-100:
Excellent answer. Almost all important information is correct
and relevant.

75-89:
Good answer. Most important information is correct but there
may be some missing details.

50-74:
Average answer. Some important information is correct but
several points are missing or incorrect.

25-49:
Weak answer. Only a small amount of relevant information is
correct.

0-24:
Very poor answer. The answer is mostly incorrect, irrelevant,
or missing.


Consider:

1. Correctness
2. Relevance
3. Important information included
4. Missing information
5. Incorrect information
6. Overall quality


Return ONLY valid JSON.

The JSON must follow this schema:

{match_schema}

The response must contain the word JSON.

Return:

score:
A number from 0 to 100.

details:
A concise explanation of why the answer received this score.
"""

    messages = [
        {
            "role": "system",
            "content": "You are an expert answer evaluator.\n\nReturn ONLY valid JSON.\nThe response must contain the word JSON.",
        },
        {"role": "user", "content": prompt},
    ]

    response = client.chat.completions.create(
        model=model,
        messages=messages,
        response_format={"type": "json_object"},
    )

    raw_output = response.choices[0].message.content
    data = json.loads(raw_output)

    # `details` sometimes comes back as a dict from the LLM; normalize to string
    if isinstance(data.get("details"), dict):
        data["details"] = json.dumps(data["details"])

    return MatchResult(**data)


# ============================================================
# SESSION STATE INIT
# ============================================================

if "baseline" not in st.session_state:
    st.session_state.baseline = None  # will hold a QA object

if "results" not in st.session_state:
    st.session_state.results = []  # list of dicts


# ============================================================
# SIDEBAR - SETTINGS
# ============================================================

st.sidebar.header("⚙️ Settings")

default_key = os.getenv("GROQ_API_KEY", "")
api_key = st.sidebar.text_input(
    "Groq API Key",
    value=default_key,
    type="password",
    help="Loaded from .env if present. You can override it here.",
)

model = st.sidebar.text_input("Model", value="openai/gpt-oss-120b")

delay_seconds = st.sidebar.slider(
    "Delay between API calls (seconds)",
    min_value=0,
    max_value=10,
    value=2,
    help="Small delay to avoid rate limiting.",
)

st.sidebar.markdown("---")

use_semantic_similarity = st.sidebar.checkbox(
    "Also compute embedding-based semantic similarity",
    value=SEMANTIC_MODEL_AVAILABLE,
    disabled=not SEMANTIC_MODEL_AVAILABLE,
    help=(
        "Adds a second, numeric similarity score (0-100) based on sentence "
        "embeddings, independent of the LLM's judgment. Useful as a sanity "
        "check since it catches meaning-based closeness even with very "
        "different wording."
    ),
)

if not SEMANTIC_MODEL_AVAILABLE:
    st.sidebar.caption(
        "⚠️ `sentence-transformers` not installed — run "
        "`pip install sentence-transformers` to enable this."
    )

st.sidebar.markdown("---")
st.sidebar.caption("Built for evaluating individual answers against a baseline reference answer.")


# ============================================================
# MAIN TITLE
# ============================================================

st.title("📝 Q&A Evaluation App")
st.write(
    "Provide a **baseline question + answer**, then upload **individual answer "
    "files** (PDF or DOCX) to automatically extract, compare, and score each one."
)


# ============================================================
# STEP 1: BASELINE QUESTION + ANSWER
# ============================================================

st.header("1️⃣ Baseline Question & Answer")

baseline_mode = st.radio(
    "How would you like to provide the baseline?",
    ["Upload a file", "Enter manually"],
    horizontal=True,
)

if baseline_mode == "Upload a file":
    baseline_file = st.file_uploader(
        "Upload baseline question + answer file (PDF or DOCX)",
        type=["pdf", "docx"],
        key="baseline_uploader",
    )

    if baseline_file is not None:
        if st.button("Extract baseline Q&A from file"):
            if not api_key:
                st.error("Please provide a Groq API key in the sidebar first.")
            else:
                with st.spinner("Reading file and extracting Question & Answer..."):
                    try:
                        text = read_uploaded_file(baseline_file)
                        if not text.strip():
                            st.error("Could not extract any text from the file.")
                        else:
                            client = get_client(api_key)
                            parsed = parse_QA(client, model, text)
                            st.session_state.baseline = QA(
                                Question=parsed.Ques, Answer=parsed.Answer
                            )
                            st.success("Baseline extracted successfully.")
                    except Exception as e:
                        st.error(f"Error extracting baseline: {e}")

    # Allow manual edit/confirmation of the extracted baseline
    if st.session_state.baseline is not None:
        st.text_input(
            "Baseline Question",
            value=st.session_state.baseline.Question,
            key="baseline_q_edit",
        )
        st.text_area(
            "Baseline Answer",
            value=st.session_state.baseline.Answer,
            key="baseline_a_edit",
            height=150,
        )
        # keep session_state.baseline synced with any manual edits
        st.session_state.baseline = QA(
            Question=st.session_state.baseline_q_edit,
            Answer=st.session_state.baseline_a_edit,
        )

else:
    manual_q = st.text_input("Baseline Question")
    manual_a = st.text_area("Baseline Answer", height=150)

    if manual_q.strip() and manual_a.strip():
        st.session_state.baseline = QA(Question=manual_q, Answer=manual_a)


if st.session_state.baseline is not None:
    with st.expander("✅ Current baseline (click to view)"):
        st.markdown(f"**Question:** {st.session_state.baseline.Question}")
        st.markdown(f"**Answer:**\n\n{st.session_state.baseline.Answer}")


# ============================================================
# STEP 2: INDIVIDUAL ANSWER FILES
# ============================================================

st.header("2️⃣ Individual Answer Files")

individual_files = st.file_uploader(
    "Upload one or more individual answer files (PDF or DOCX)",
    type=["pdf", "docx"],
    accept_multiple_files=True,
    key="individual_uploader",
)

if individual_files:
    st.info(f"{len(individual_files)} file(s) ready to be evaluated.")


# ============================================================
# STEP 3: RUN EVALUATION
# ============================================================

st.header("3️⃣ Run Evaluation")

run_disabled = (
    st.session_state.baseline is None
    or not individual_files
    or not api_key
)

if st.session_state.baseline is None:
    st.warning("Set a baseline question & answer first.")
if not individual_files:
    st.warning("Upload at least one individual answer file.")
if not api_key:
    st.warning("Enter your Groq API key in the sidebar.")

if st.button("🚀 Run Evaluation", disabled=run_disabled, type="primary"):
    client = get_client(api_key)
    results = []
    progress = st.progress(0.0)
    status = st.empty()

    total = len(individual_files)

    semantic_model = None
    if use_semantic_similarity and SEMANTIC_MODEL_AVAILABLE:
        with st.spinner("Loading semantic similarity model..."):
            semantic_model = load_semantic_model()

    for i, uploaded_file in enumerate(individual_files):
        status.write(f"Processing **{uploaded_file.name}** ({i + 1}/{total})...")

        try:
            text = read_uploaded_file(uploaded_file)

            if not text.strip():
                results.append(
                    {
                        "file": uploaded_file.name,
                        "question": "",
                        "answer": "",
                        "score": None,
                        "semantic_similarity": None,
                        "details": "Could not extract text from file.",
                    }
                )
                progress.progress((i + 1) / total)
                continue

            # LLM call 1: extract individual Q&A
            parsed_individual = parse_QA(client, model, text)

            time.sleep(delay_seconds)

            # LLM call 2: score against baseline (meaning-based grading)
            result = final_score(client, model, st.session_state.baseline, parsed_individual)

            # Optional: embedding-based semantic similarity (independent signal)
            sim_score = None
            if semantic_model is not None:
                sim_score = semantic_similarity(
                    semantic_model,
                    st.session_state.baseline.Answer,
                    parsed_individual.Answer,
                )

            results.append(
                {
                    "file": uploaded_file.name,
                    "question": parsed_individual.Ques,
                    "answer": parsed_individual.Answer,
                    "score": result.score,
                    "semantic_similarity": sim_score,
                    "details": result.details,
                }
            )

            time.sleep(delay_seconds)

        except Exception as e:
            results.append(
                {
                    "file": uploaded_file.name,
                    "question": "",
                    "answer": "",
                    "score": None,
                    "semantic_similarity": None,
                    "details": f"Error: {e}",
                }
            )

        progress.progress((i + 1) / total)

    status.write("Done ✅")
    st.session_state.results = results


# ============================================================
# STEP 4: RESULTS
# ============================================================

if st.session_state.results:
    st.header("4️⃣ Results")

    valid_results = [r for r in st.session_state.results if r["score"] is not None]
    failed_results = [r for r in st.session_state.results if r["score"] is None]

    valid_results.sort(key=lambda r: r["score"], reverse=True)

    df = pd.DataFrame(valid_results)

    if not df.empty:
        st.subheader("📊 All Results (sorted by score)")

        display_cols = ["file", "question", "score"]
        if "semantic_similarity" in df.columns and df["semantic_similarity"].notna().any():
            display_cols.append("semantic_similarity")
        display_cols.append("details")

        st.dataframe(
            df[display_cols].rename(
                columns={"semantic_similarity": "semantic similarity %"}
            ),
            use_container_width=True,
            hide_index=True,
        )

        if "semantic_similarity" in df.columns and df["semantic_similarity"].notna().any():
            st.caption(
                "**score** = LLM's holistic judgment (correctness, relevance, "
                "completeness). **semantic similarity %** = embedding-based "
                "cosine similarity between the reference answer and the "
                "individual's answer — a numeric check of meaning-closeness, "
                "independent of the LLM."
            )

        csv = df.to_csv(index=False).encode("utf-8")
        st.download_button(
            "⬇️ Download results as CSV",
            data=csv,
            file_name="qa_evaluation_results.csv",
            mime="text/csv",
        )

        col1, col2 = st.columns(2)

        with col1:
            st.subheader("🏆 Top 2 Candidates")
            for rank, candidate in enumerate(df.head(2).to_dict("records"), start=1):
                st.markdown(f"**{rank}. {candidate['file']}** — {candidate['score']}%")
                st.caption(candidate["details"])

        with col2:
            st.subheader("⚠️ Lowest 2 Candidates")
            for rank, candidate in enumerate(df.tail(2).to_dict("records"), start=1):
                st.markdown(f"**{rank}. {candidate['file']}** — {candidate['score']}%")
                st.caption(candidate["details"])

        with st.expander("View extracted individual answers"):
            for r in valid_results:
                st.markdown(f"**{r['file']}**")
                st.markdown(f"*Question:* {r['question']}")
                st.markdown(f"*Answer:* {r['answer']}")
                st.markdown("---")

    if failed_results:
        st.subheader("❌ Files that failed to process")
        for r in failed_results:
            st.markdown(f"- **{r['file']}**: {r['details']}")
