"""
extract_worksheet_qa.py

1. Sends a photo of a worksheet (handwritten or printed) to Claude's
   vision model.
2. Claude returns structured JSON: every question, its type, and its
   correct answer.
3. The script writes a Word document containing ONLY the questions
   (blanks restored, answers removed) so you can print it fresh for
   the next practice round.

------------------------------------------------------------------
SETUP
------------------------------------------------------------------
  pip install anthropic python-docx

  Get an API key from https://console.anthropic.com  (Settings > API Keys)
  Set it as an environment variable so it's not hard-coded in the script:

  Windows (PowerShell):
      setx ANTHROPIC_API_KEY "your-key-here"
      (then close and reopen PowerShell so it picks up the new variable)

  Mac/Linux:
      export ANTHROPIC_API_KEY="your-key-here"

------------------------------------------------------------------
RUN
------------------------------------------------------------------
  python extract_worksheet_qa.py path\to\worksheet_photo.jpg

  Output:
    - qa_extracted.json          (questions + answers, for your records)
    - worksheet_questions.docx   (questions only, ready to print/reuse)
------------------------------------------------------------------
"""

import sys
import os
import json
import base64
import mimetypes

import anthropic
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


EXTRACTION_PROMPT = """You are reading a photo of a kids' worksheet (it may be
handwritten and messy). Extract every question into this exact JSON shape
and return ONLY the JSON, nothing else:

{
  "title": "short worksheet title",
  "sections": [
    {
      "section_name": "Fill in the Blanks",
      "type": "fill_blank",
      "questions": [
        {"prompt": "Market has many ______.", "answer": "shops"}
      ]
    },
    {
      "section_name": "Tick the Correct Option",
      "type": "multiple_choice",
      "questions": [
        {"prompt": "We call ______ when there is a fire.",
         "options": ["Police station", "Bank", "Fire station"],
         "answer": "Fire station"}
      ]
    },
    {
      "section_name": "True or False",
      "type": "true_false",
      "questions": [
        {"prompt": "We go to school to study and learn.", "answer": true}
      ]
    },
    {
      "section_name": "Missing Letter",
      "type": "missing_letter",
      "questions": [
        {"prompt": "S_HOOL", "answer": "SCHOOL"}
      ]
    },
    {
      "section_name": "Re-arrange the Letters",
      "type": "rearrange",
      "questions": [
        {"prompt": "SOHPS", "answer": "SHOPS"}
      ]
    },
    {
      "section_name": "Answer the Following",
      "type": "short_answer",
      "questions": [
        {"prompt": "Where do we go to post letters or parcels?", "answer": "Post office"}
      ]
    }
  ]
}

Only include section types that actually appear in the image. Keep the
original wording of each question. If a blank was already filled in with
an answer, put the filled word in "answer" and restore the blank (______)
in "prompt"."""


def encode_image(path):
    mime = mimetypes.guess_type(path)[0] or "image/jpeg"
    with open(path, "rb") as f:
        data = base64.standard_b64encode(f.read()).decode("utf-8")
    return mime, data


def extract_questions_and_answers(image_path):
    client = anthropic.Anthropic()  # reads ANTHROPIC_API_KEY from environment
    mime, image_b64 = encode_image(image_path)

    response = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=2000,
        messages=[{
            "role": "user",
            "content": [
                {"type": "image", "source": {"type": "base64", "media_type": mime, "data": image_b64}},
                {"type": "text", "text": EXTRACTION_PROMPT},
            ],
        }],
    )

    raw_text = response.content[0].text.strip()
    raw_text = raw_text.removeprefix("```json").removeprefix("```").removesuffix("```").strip()
    return json.loads(raw_text)


def build_questions_doc(data, output_path):
    doc = Document()

    title = doc.add_heading(data.get("title", "Practice Worksheet"), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    name_line = doc.add_paragraph("Name: ______________________     Date: ____________")
    name_line.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for section in data["sections"]:
        doc.add_heading(section["section_name"], level=2)
        qtype = section["type"]

        for i, q in enumerate(section["questions"], start=1):
            if qtype == "multiple_choice":
                doc.add_paragraph(f"{i}. {q['prompt']}")
                for opt in q["options"]:
                    doc.add_paragraph(opt, style="List Bullet")
            elif qtype == "true_false":
                doc.add_paragraph(f"{i}. {q['prompt']}   ___________")
            elif qtype in ("missing_letter", "rearrange"):
                doc.add_paragraph(f"{i}. {q['prompt']}   =   ______________")
            elif qtype == "short_answer":
                doc.add_paragraph(f"{i}. {q['prompt']}")
                doc.add_paragraph("_" * 45)
            else:  # fill_blank and anything else
                doc.add_paragraph(f"{i}. {q['prompt']}")

    doc.save(output_path)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python extract_worksheet_qa.py path\\to\\image.jpg")
        sys.exit(1)

    image_path = sys.argv[1]
    print("Sending image to Claude for extraction...")
    data = extract_questions_and_answers(image_path)

    with open("qa_extracted.json", "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2)
    print("Saved questions + answers to qa_extracted.json")

    build_questions_doc(data, "worksheet_questions.docx")
    print("Saved questions-only worksheet to worksheet_questions.docx")
